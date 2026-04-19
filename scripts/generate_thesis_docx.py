"""
根据同济大学毕业设计（论文）模板 2026 版格式规范，
将 thesis/ 目录下的 Markdown 内容生成为规范的 .docx 文件。

格式规范（依据模板气泡标注）：
- 封面课题：小二号黑体居中
- 章标题（1级）：小三号 黑体 居中 每章换页+空一行 段前后 0.5 行 行距 1.5 倍
- 节标题（2级）：小四号 黑体 顶格 段前后 0.5 行 行距 1.5 倍
- 小节标题（3级）：小四号 黑体 缩进 2 汉字 序号与题名空一格
- 三级以下标题：小四号 黑体 缩进 2 汉字，使用 (1)/(2) 或 ①/② 编号
- 正文：小四号 宋体/TNR 两端对齐 首行缩进 2 汉字 行距 1.5 倍
- 摘要/谢辞标题：四号 黑体 居中
- 目录标题：小三号 黑体 居中（"目  录"中间空一格）
- 英文摘要课题：小二号 TNR 加粗居中
- ABSTRACT：四号 TNR 加粗居中
- 表题：小五号 宋体 加粗 居中 表上方；表序"表 X.Y"
- 图题：小五号 宋体 加粗 居中 图下方；图序"图 X.Y"
- 公式：居中，编号 (X.Y) 右对齐
- 参考文献：顶格 悬挂缩进 小四
- 页眉："毕业设计（论文）"居中小五；页脚页码居中
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 字号常量 ──
CHU_HAO = Pt(36)     # 初号
XIAO_CHU = Pt(28)    # 小初
YI_HAO = Pt(26)      # 一号
XIAO_YI = Pt(24)     # 小一
ER_HAO = Pt(22)      # 二号
XIAO_ER = Pt(18)     # 小二
SAN_HAO = Pt(16)     # 三号
XIAO_SAN = Pt(15)    # 小三
SI_HAO = Pt(14)      # 四号
XIAO_SI = Pt(12)     # 小四
WU_HAO = Pt(10.5)    # 五号
XIAO_WU = Pt(9)      # 小五

# ── 字体常量 ──
HEITI = 'SimHei'
SONGTI = 'SimSun'
TNR = 'Times New Roman'
KAITI = 'KaiTi'
CODE_FONT = 'Consolas'

CHAPTER_MAP = {
    '第一章 引言': '1  引言',
    '第二章 相关技术': '2  相关技术',
    '第三章 系统需求分析与设计': '3  系统需求分析与设计',
    '第四章 系统实现': '4  系统实现',
    '第五章 系统优化': '5  系统优化',
    '第六章 系统测试与结果分析': '6  系统测试与结果分析',
    '第七章 结论与展望': '7  结论与展望',
}

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
THESIS_DIR = os.path.join(BASE_DIR, 'thesis')
FIG_DIR = os.path.join(THESIS_DIR, 'figures')


# ── 底层样式工具 ──

def set_run_font(run, font_name=TNR, east_asia=SONGTI, size=XIAO_SI,
                 bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:eastAsia="{east_asia}" '
            f'w:ascii="{font_name}" w:hAnsi="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), east_asia)
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
    if color:
        run.font.color.rgb = color


def set_paragraph_format(para, alignment=None, first_indent=None,
                         left_indent=None, right_indent=None,
                         space_before=None, space_after=None,
                         line_spacing=None, line_spacing_rule=None,
                         keep_with_next=False, page_break_before=False):
    pf = para.paragraph_format
    if alignment is not None:
        pf.alignment = alignment
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if right_indent is not None:
        pf.right_indent = right_indent
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if line_spacing_rule is not None:
        pf.line_spacing_rule = line_spacing_rule
    if keep_with_next:
        pf.keep_with_next = True
    if page_break_before:
        pf.page_break_before = True


def add_formatted_paragraph(doc, text, font_name=TNR, east_asia=SONGTI,
                            size=XIAO_SI, bold=False, alignment=None,
                            first_indent=None, space_before=None,
                            space_after=None, line_spacing=1.5,
                            page_break_before=False, keep_with_next=False):
    para = doc.add_paragraph()
    set_paragraph_format(
        para, alignment=alignment, first_indent=first_indent,
        space_before=space_before, space_after=space_after,
        line_spacing=line_spacing, line_spacing_rule=WD_LINE_SPACING.MULTIPLE,
        page_break_before=page_break_before, keep_with_next=keep_with_next)
    run = para.add_run(text)
    set_run_font(run, font_name=font_name, east_asia=east_asia,
                 size=size, bold=bold)
    return para


def add_empty_paragraph(doc, size=XIAO_SI):
    """添加一个空段（用于"空一行"）"""
    para = doc.add_paragraph()
    set_paragraph_format(
        para, line_spacing=1.5, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    run = para.add_run('')
    set_run_font(run, east_asia=SONGTI, size=size)
    return para


def add_rich_paragraph(doc, text, default_size=XIAO_SI, default_east_asia=SONGTI,
                       default_font=TNR, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       first_indent=Cm(0.74),
                       space_before=None, space_after=None, line_spacing=1.5):
    """支持 **加粗** 和 `代码` 内联的段落"""
    para = doc.add_paragraph()
    set_paragraph_format(
        para, alignment=alignment, first_indent=first_indent,
        space_before=space_before, space_after=space_after,
        line_spacing=line_spacing, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    pattern = r'(\*\*(.+?)\*\*|`(.+?)`)'
    last_end = 0
    for m in re.finditer(pattern, text):
        if m.start() > last_end:
            run = para.add_run(text[last_end:m.start()])
            set_run_font(run, font_name=default_font,
                         east_asia=default_east_asia, size=default_size)
        if m.group(2):
            run = para.add_run(m.group(2))
            set_run_font(run, font_name=default_font,
                         east_asia=default_east_asia, size=default_size,
                         bold=True)
        elif m.group(3):
            run = para.add_run(m.group(3))
            set_run_font(run, font_name=CODE_FONT, east_asia=CODE_FONT,
                         size=default_size)
        last_end = m.end()
    if last_end < len(text):
        run = para.add_run(text[last_end:])
        set_run_font(run, font_name=default_font,
                     east_asia=default_east_asia, size=default_size)
    return para


# ── 章节标题样式 ──

def add_chapter_heading(doc, title, page_break=True):
    """1 级标题：小三 黑体 居中 换页+空一行 段前后 0.5 行"""
    # 章前先换页（page_break_before 在段落自身），并空一行
    if page_break:
        # 空段落承载 page break
        p_break = doc.add_paragraph()
        set_paragraph_format(p_break, page_break_before=True,
                             line_spacing=1.5,
                             line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
        run = p_break.add_run('')
        set_run_font(run, east_asia=SONGTI, size=XIAO_SI)

    return add_formatted_paragraph(
        doc, title, font_name=TNR, east_asia=HEITI,
        size=XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(12),
        line_spacing=1.5,
        keep_with_next=True)


def add_section_heading(doc, title):
    """2 级标题：小四 黑体 顶格 段前后 0.5 行"""
    return add_formatted_paragraph(
        doc, title, font_name=TNR, east_asia=HEITI,
        size=XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=None,
        space_before=Pt(12), space_after=Pt(12),
        line_spacing=1.5,
        keep_with_next=True)


def add_subsection_heading(doc, title):
    """3 级标题：小四 黑体 缩进 2 汉字 段前后 0.5 行"""
    return add_formatted_paragraph(
        doc, title, font_name=TNR, east_asia=HEITI,
        size=XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=Cm(0.74),
        space_before=Pt(12), space_after=Pt(12),
        line_spacing=1.5,
        keep_with_next=True)


def add_sub_sub_heading(doc, title):
    """三级以下占行标题：小四 黑体 缩进 2 汉字"""
    return add_formatted_paragraph(
        doc, title, font_name=TNR, east_asia=HEITI,
        size=XIAO_SI, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=Cm(0.74),
        space_before=Pt(6), space_after=Pt(6),
        line_spacing=1.5,
        keep_with_next=True)


def add_body_paragraph(doc, text):
    """正文段落：小四 宋体/TNR 两端对齐 首行缩进 2 汉字 行距 1.5"""
    if not text.strip():
        return None
    return add_rich_paragraph(
        doc, text, default_size=XIAO_SI, default_east_asia=SONGTI,
        default_font=TNR, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_indent=Cm(0.74), line_spacing=1.5)


def add_code_block(doc, code_text):
    """代码块：五号 Consolas 单倍行距"""
    for line in code_text.split('\n'):
        para = doc.add_paragraph()
        set_paragraph_format(
            para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_indent=Cm(0.74),
            space_before=Pt(0), space_after=Pt(0),
            line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
        run = para.add_run(line)
        set_run_font(run, font_name=CODE_FONT, east_asia=CODE_FONT,
                     size=WU_HAO)
        # 浅灰底纹
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            shd = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
            rPr.append(shd)


# ── 表格 ──

def add_table_caption(doc, caption):
    """表题：小五号 宋体 加粗 居中 置于表上方"""
    cap_para = doc.add_paragraph()
    set_paragraph_format(
        cap_para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(6), space_after=Pt(3),
        line_spacing=1.5, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    run = cap_para.add_run(caption)
    set_run_font(run, east_asia=SONGTI, size=XIAO_WU, bold=True)


def add_table_from_md(doc, header_line, rows, caption=None):
    headers = [c.strip() for c in header_line.strip('|').split('|')]
    data_rows = []
    for row in rows:
        cells = [c.strip() for c in row.strip('|').split('|')]
        data_rows.append(cells)
    ncols = len(headers)

    if caption:
        add_table_caption(doc, caption)

    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        para = cell.paragraphs[0]
        set_paragraph_format(
            para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
        run = para.add_run(h)
        set_run_font(run, east_asia=SONGTI, size=XIAO_WU, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row_data in enumerate(data_rows):
        for j in range(min(len(row_data), ncols)):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            set_paragraph_format(
                para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
            run = para.add_run(row_data[j])
            set_run_font(run, east_asia=SONGTI, size=XIAO_WU)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _apply_three_line_table(table)
    doc.add_paragraph()


def _apply_three_line_table(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
        <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}></w:tcPr>')
                tc.insert(0, tcPr)
            tcPr.append(parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="0" '
                f'w:color="000000"/></w:tcBorders>'))


# ── 图片 ──

def add_figure(doc, img_path, caption, width_cm=14):
    """插入图片。
    图居中，上下各空一行；图题小五号宋体加粗居中置于图下方"""
    # 上方空行
    add_empty_paragraph(doc, size=XIAO_SI)

    # 图片段
    para = doc.add_paragraph()
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         line_spacing=1.0,
                         line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    run = para.add_run()
    try:
        run.add_picture(img_path, width=Cm(width_cm))
    except FileNotFoundError:
        run = para.add_run(f'[缺失图片：{img_path}]')
        set_run_font(run, east_asia=SONGTI, size=XIAO_SI)

    # 图题
    cap_para = doc.add_paragraph()
    set_paragraph_format(cap_para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(3), space_after=Pt(6),
                         line_spacing=1.5,
                         line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    run = cap_para.add_run(caption)
    set_run_font(run, east_asia=SONGTI, size=XIAO_WU, bold=True)

    # 下方空行
    add_empty_paragraph(doc, size=XIAO_SI)


# ── 公式 ──

def add_formula(doc, formula_text, number=None):
    """公式：居中小四 TNR 斜体；如提供 number 则编号 (X.Y) 右对齐"""
    para = doc.add_paragraph()
    set_paragraph_format(
        para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(6), space_after=Pt(6),
        line_spacing=1.5, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)

    if number:
        # 使用 tab stop 实现 公式居中 + 编号右对齐
        # 左/中/右三 tab stop
        pf = para.paragraph_format
        tab_stops = pf.tab_stops
        tab_stops.add_tab_stop(Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
        tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)

        run = para.add_run('\t')
        set_run_font(run, font_name=TNR, east_asia=SONGTI, size=XIAO_SI)
        run = para.add_run(formula_text)
        set_run_font(run, font_name=TNR, east_asia=SONGTI, size=XIAO_SI,
                     italic=True)
        run = para.add_run(f'\t({number})')
        set_run_font(run, font_name=TNR, east_asia=SONGTI, size=XIAO_SI)
    else:
        run = para.add_run(formula_text)
        set_run_font(run, font_name=TNR, east_asia=SONGTI, size=XIAO_SI,
                     italic=True)


# ── 封面 ──

def add_cover_page(doc):
    para = doc.add_paragraph()
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(60), space_after=Pt(6))
    run = para.add_run('TONGJI UNIVERSITY')
    set_run_font(run, font_name=TNR, east_asia=TNR, size=Pt(26), bold=True)

    para = doc.add_paragraph()
    set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=Pt(12), space_after=Pt(36))
    run = para.add_run('毕业设计（论文）')
    set_run_font(run, font_name=TNR, east_asia=HEITI, size=Pt(36), bold=True)

    # 课题名称：小二号 黑体 加粗 居中
    para = doc.add_paragraph()
    set_paragraph_format(
        para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(24), space_after=Pt(0),
        line_spacing=1.5, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    run = para.add_run('课题名称       面向知识增强的语音问答')
    set_run_font(run, east_asia=HEITI, size=XIAO_ER, bold=True)

    para = doc.add_paragraph()
    set_paragraph_format(
        para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(0), space_after=Pt(24),
        line_spacing=1.5, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    run = para.add_run('Agent 设计与优化')
    set_run_font(run, east_asia=HEITI, size=XIAO_ER, bold=True)

    info_items = [
        ('学    院', '计算机科学与技术学院'),
        ('专　　业', '软件工程'),
        ('学生姓名', '黄志栋'),
        ('学　　号', '2251760'),
        ('指导教师', ''),
    ]
    for label, value in info_items:
        para = doc.add_paragraph()
        set_paragraph_format(
            para, alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=Pt(3), space_after=Pt(3),
            line_spacing=1.5, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
        run = para.add_run(f'{label}　　{value}')
        set_run_font(run, east_asia=SONGTI, size=SI_HAO)

    doc.add_page_break()


# ── 摘要 ──

def add_chinese_abstract(doc, title, abstract_text, keywords):
    # 课题名称：小二号 黑体 加粗 居中 段前后 0.5 行，题目前后各空一行
    add_empty_paragraph(doc)
    add_formatted_paragraph(
        doc, title, east_asia=HEITI, size=XIAO_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(12), line_spacing=1.5)
    add_empty_paragraph(doc)

    # "摘  要"：四号 黑体 加粗 居中
    add_formatted_paragraph(
        doc, '摘　要', east_asia=HEITI, size=SI_HAO, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(12), line_spacing=1.5)

    for p in abstract_text.strip().split('\n\n'):
        p = p.strip()
        if p:
            add_body_paragraph(doc, p)

    # 关键词：小四 宋体加粗顶格 + 小四 宋体
    para = doc.add_paragraph()
    set_paragraph_format(
        para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=None,
        space_before=Pt(12), space_after=Pt(0),
        line_spacing=1.5, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    run = para.add_run('关键词：')
    set_run_font(run, east_asia=SONGTI, size=XIAO_SI, bold=True)
    # 关键词中文逗号分隔，最后无标点
    kws = re.split(r'[；;，,、]', keywords)
    kws = [k.strip() for k in kws if k.strip()]
    run = para.add_run('，'.join(kws))
    set_run_font(run, east_asia=SONGTI, size=XIAO_SI)
    doc.add_page_break()


def add_english_abstract(doc, title, abstract_text, keywords):
    # 英文课题：小二号 TNR 加粗 居中 首字母大写 段前后 0.5 行；题目前后各空一行
    add_empty_paragraph(doc)
    add_formatted_paragraph(
        doc, title, font_name=TNR, east_asia=TNR, size=XIAO_ER, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(12), line_spacing=1.5)
    add_empty_paragraph(doc)

    # ABSTRACT：四号 TNR 加粗 居中
    add_formatted_paragraph(
        doc, 'ABSTRACT', font_name=TNR, east_asia=TNR, size=SI_HAO, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(12), line_spacing=1.5)

    # 正文：小四 TNR 首行缩进 2 汉字
    for p in abstract_text.strip().split('\n\n'):
        p = p.strip()
        if p:
            add_rich_paragraph(
                doc, p, default_size=XIAO_SI, default_east_asia=TNR,
                default_font=TNR,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_indent=Cm(0.74), line_spacing=1.5)

    # Keywords: 小四 TNR 加粗顶格；关键词逗号+空格分隔
    para = doc.add_paragraph()
    set_paragraph_format(
        para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent=None,
        space_before=Pt(12), space_after=Pt(0),
        line_spacing=1.5, line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    run = para.add_run('Keywords: ')
    set_run_font(run, font_name=TNR, east_asia=TNR, size=XIAO_SI, bold=True)
    kws = re.split(r'[;；,，]', keywords)
    kws = [k.strip() for k in kws if k.strip()]
    run = para.add_run(', '.join(kws))
    set_run_font(run, font_name=TNR, east_asia=TNR, size=XIAO_SI)
    doc.add_page_break()


def add_toc_placeholder(doc):
    """目  录：小三号 黑体 居中 段前后 0.5 行 行距 1.5"""
    add_formatted_paragraph(
        doc, '目　录', east_asia=HEITI, size=XIAO_SAN, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(24), space_after=Pt(18), line_spacing=1.5)
    para = doc.add_paragraph()
    fld_char_begin = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    fld_code = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> '
        f'TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    fld_char_sep = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fld_char_end = parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run1 = para.add_run()
    run1._element.append(fld_char_begin)
    run2 = para.add_run()
    run2._element.append(fld_code)
    run3 = para.add_run()
    run3._element.append(fld_char_sep)
    run4 = para.add_run('（请在 Word 中右键点击此处，选择"更新域"以生成目录）')
    set_run_font(run4, east_asia=SONGTI, size=XIAO_SI)
    run5 = para.add_run()
    run5._element.append(fld_char_end)
    doc.add_page_break()


# ── Markdown 解析 ──

def parse_abstract_md(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    cn_abstract_match = re.search(
        r'## 摘要\s*\n(.*?)(?=\*\*关键词\*\*)', content, re.DOTALL)
    cn_abstract = cn_abstract_match.group(1).strip() if cn_abstract_match else ''
    cn_kw_match = re.search(r'\*\*关键词\*\*：(.+)', content)
    cn_keywords = cn_kw_match.group(1).strip() if cn_kw_match else ''
    en_abstract_match = re.search(
        r'## Abstract\s*\n(.*?)(?=\*\*Keywords\*\*)', content, re.DOTALL)
    en_abstract = en_abstract_match.group(1).strip() if en_abstract_match else ''
    en_kw_match = re.search(r'\*\*Keywords\*\*:\s*(.+)', content)
    en_keywords = en_kw_match.group(1).strip() if en_kw_match else ''
    return cn_abstract, cn_keywords, en_abstract, en_keywords


# 正则：Markdown 图片 ![caption](path)
IMG_PATTERN = re.compile(r'!\[([^\]]*)\]\(([^)]+)\)')
# 表标题：紧邻表格前的段落中以 "表 X.Y" 或 "表X.Y" 开头
TABLE_CAP_PATTERN = re.compile(r'(表\s*\d+[\.\-]\d+[^\n]*)')
# 公式编号：$$...$$ 形式，行尾可能含 (X.Y)
FORMULA_NUM_PATTERN = re.compile(r'\((\d+\.\d+)\)\s*$')


def render_chapter_md(doc, filepath, chapter_num):
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    prev_text = ''
    skip_references = False

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # 过滤"## 参考文献" "## 致谢" 这两节（在第七章末尾，转由独立函数渲染）
        if line.startswith('## 参考文献') or line.startswith('## 致谢'):
            skip_references = True
            i += 1
            continue
        if skip_references:
            i += 1
            continue

        # 代码块
        if line.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # 图片：![caption](path)
        img_match = IMG_PATTERN.match(line.strip())
        if img_match:
            caption = img_match.group(1)
            img_rel_path = img_match.group(2)
            img_path = os.path.join(THESIS_DIR, img_rel_path)
            add_figure(doc, img_path, caption)
            i += 1
            continue

        # 一级标题
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            mapped_title = CHAPTER_MAP.get(title, title)
            add_chapter_heading(doc, mapped_title, page_break=True)
            i += 1
            continue

        # 二级标题
        if line.startswith('## ') and not line.startswith('### '):
            title = line[3:].strip()
            add_section_heading(doc, title)
            i += 1
            continue

        # 三级标题
        if line.startswith('### ') and not line.startswith('#### '):
            title = line[4:].strip()
            add_subsection_heading(doc, title)
            i += 1
            continue

        # 四级标题（三级以下占行标题）
        if line.startswith('#### '):
            title = line[5:].strip()
            add_sub_sub_heading(doc, title)
            i += 1
            continue

        # 公式块 $$...$$
        if line.startswith('$$'):
            formula_lines = [line[2:]]
            i += 1
            while i < len(lines) and not lines[i].strip().endswith('$$'):
                formula_lines.append(lines[i].rstrip('\n'))
                i += 1
            if i < len(lines):
                last = lines[i].rstrip('\n').rstrip('$')
                if last.strip():
                    formula_lines.append(last)
                i += 1
            formula_text = ' '.join(formula_lines).strip()
            num_match = FORMULA_NUM_PATTERN.search(formula_text)
            number = None
            if num_match:
                number = num_match.group(1)
                formula_text = FORMULA_NUM_PATTERN.sub('', formula_text).strip()
            add_formula(doc, formula_text, number=number)
            continue

        # 表格
        if '|' in line and i + 1 < len(lines) and re.match(
                r'\s*\|[\s\-:|]+\|', lines[i + 1]):
            header_line = line
            i += 2
            data_rows = []
            while (i < len(lines) and '|' in lines[i].rstrip('\n')
                   and lines[i].strip().startswith('|')):
                data_rows.append(lines[i].rstrip('\n'))
                i += 1
            # 从 prev_text 中寻找表题（支持 "表 X.Y" 或 "表 X-Y"）
            caption = None
            cap_match = TABLE_CAP_PATTERN.search(prev_text)
            if cap_match:
                caption = cap_match.group(1).strip()
                # 统一 X-Y → X.Y
                caption = re.sub(r'(表\s*\d+)-(\d+)', r'\1.\2', caption)
            add_table_from_md(doc, header_line, data_rows, caption=caption)
            continue

        # 普通段落（可能跨行）
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip('\n')
            if (not nxt.strip() or
                    nxt.startswith('#') or
                    nxt.startswith('```') or
                    nxt.startswith('$$') or
                    IMG_PATTERN.match(nxt.strip()) or
                    (re.match(r'\s*\|.*\|', nxt) and
                     i + 1 < len(lines) and
                     re.match(r'\s*\|[\s\-:|]+\|', lines[i + 1]))):
                break
            para_lines.append(nxt)
            i += 1

        text = ' '.join(para_lines).strip()
        # 把 X-Y 格式的表号/图号规范为 X.Y
        text = re.sub(r'(表\s*\d+)-(\d+)', r'\1.\2', text)
        text = re.sub(r'(图\s*\d+)-(\d+)', r'\1.\2', text)
        add_body_paragraph(doc, text)
        prev_text = text


# ── 参考文献与致谢 ──

def add_references(doc, ref_text):
    # 参考文献：四号 黑体 居中 换页
    p_break = doc.add_paragraph()
    set_paragraph_format(p_break, page_break_before=True,
                         line_spacing=1.5,
                         line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    run = p_break.add_run('')
    set_run_font(run, east_asia=SONGTI, size=XIAO_SI)

    add_formatted_paragraph(
        doc, '参考文献', east_asia=HEITI, size=SI_HAO, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(18), line_spacing=1.5)

    for line in ref_text.strip().split('\n'):
        line = line.strip()
        if not line:
            continue
        para = doc.add_paragraph()
        set_paragraph_format(
            para, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            first_indent=Cm(-0.74), left_indent=Cm(0.74),
            space_before=Pt(0), space_after=Pt(3),
            line_spacing=1.5,
            line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
        run = para.add_run(line)
        set_run_font(run, font_name=TNR, east_asia=SONGTI, size=XIAO_SI)


def add_acknowledgment(doc, text):
    # 致  谢：四号 黑体 居中 换页
    p_break = doc.add_paragraph()
    set_paragraph_format(p_break, page_break_before=True,
                         line_spacing=1.5,
                         line_spacing_rule=WD_LINE_SPACING.MULTIPLE)
    run = p_break.add_run('')
    set_run_font(run, east_asia=SONGTI, size=XIAO_SI)

    add_formatted_paragraph(
        doc, '致　谢', east_asia=HEITI, size=SI_HAO, bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=Pt(12), space_after=Pt(18), line_spacing=1.5)

    for p in text.strip().split('\n\n'):
        p = p.strip()
        if p:
            add_body_paragraph(doc, p)


# ── 页面设置 ──

def setup_page(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    font = style.font
    font.name = TNR
    font.size = XIAO_SI
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:eastAsia="{SONGTI}" '
            f'w:ascii="{TNR}" w:hAnsi="{TNR}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), SONGTI)

    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    # 页眉
    header = section.header
    header.is_linked_to_previous = False
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h_para.add_run('毕业设计（论文）')
    set_run_font(run, east_asia=SONGTI, size=XIAO_WU)

    # 页脚：居中页码
    footer = section.footer
    footer.is_linked_to_previous = False
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for xml_str in [
        f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>',
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>',
        f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>',
    ]:
        r = f_para.add_run()
        r._element.append(parse_xml(xml_str))
    r = f_para.add_run()
    set_run_font(r, east_asia=SONGTI, size=XIAO_WU)
    r = f_para.add_run()
    r._element.append(parse_xml(
        f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


def extract_references_and_acknowledgment(chapter7_path):
    with open(chapter7_path, 'r', encoding='utf-8') as f:
        content = f.read()
    ref_match = re.search(
        r'## 参考文献\s*\n(.*?)(?=## 致谢)', content, re.DOTALL)
    ref_text = ref_match.group(1).strip() if ref_match else ''
    ack_match = re.search(r'## 致谢\s*\n(.*?)$', content, re.DOTALL)
    ack_text = ack_match.group(1).strip() if ack_match else ''
    return ref_text, ack_text


def main():
    doc = Document()
    setup_page(doc)

    add_cover_page(doc)

    abstract_path = os.path.join(THESIS_DIR, '00_abstract.md')
    cn_abstract, cn_keywords, en_abstract, en_keywords = parse_abstract_md(
        abstract_path)

    cn_title = '面向知识增强的语音问答 Agent 设计与优化'
    add_chinese_abstract(doc, cn_title, cn_abstract, cn_keywords)

    en_title = ('Design and Optimization of Knowledge-Enhanced '
                'Voice Question-Answering Agent')
    add_english_abstract(doc, en_title, en_abstract, en_keywords)

    add_toc_placeholder(doc)

    chapters = [
        ('chapter1_introduction.md', 1),
        ('chapter2_technologies.md', 2),
        ('chapter3_design.md', 3),
        ('chapter4_implementation.md', 4),
        ('chapter5_optimization.md', 5),
        ('chapter6_testing.md', 6),
        ('chapter7_conclusion.md', 7),
    ]
    for filename, ch_num in chapters:
        filepath = os.path.join(THESIS_DIR, filename)
        if os.path.exists(filepath):
            render_chapter_md(doc, filepath, ch_num)

    ch7_path = os.path.join(THESIS_DIR, 'chapter7_conclusion.md')
    ref_text, ack_text = extract_references_and_acknowledgment(ch7_path)
    if ref_text:
        add_references(doc, ref_text)
    if ack_text:
        add_acknowledgment(doc, ack_text)

    output_path = os.path.join(BASE_DIR, '毕业论文.docx')
    doc.save(output_path)
    print(f'论文已生成：{output_path}')


if __name__ == '__main__':
    main()
